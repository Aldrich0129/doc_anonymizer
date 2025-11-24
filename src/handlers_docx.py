"""Manejador sencillo para DOCX con reemplazos basados en reglas."""
from pathlib import Path
from typing import Iterable

from docx import Document

from anonymizer_core import anonymize_text


def _iter_all_paragraphs(doc: Document) -> Iterable:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def anonymize_docx(input_path: str, output_path: str, config_path: str):
    """Carga un DOCX, aplica reglas y guarda el resultado.

    Esta implementación prioriza la simplicidad: reemplaza el texto completo
    de cada párrafo, por lo que puede perderse parte del formato si el texto
    original estaba dividido en varios *runs*.
    """
    doc = Document(input_path)

    for paragraph in _iter_all_paragraphs(doc):
        new_text, _ = anonymize_text(paragraph.text, config_path)
        paragraph.text = new_text

    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
