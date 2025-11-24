# src/handlers_pdf.py
from pathlib import Path
from io import BytesIO
import warnings
import xml.etree.ElementTree as ET
import shutil
from cryptography.utils import CryptographyDeprecationWarning

# 避免 pypdf 在导入时因弃用 ARC4 发出的噪声告警
warnings.filterwarnings(
    "ignore",
    category=CryptographyDeprecationWarning,
    module="pypdf\\._crypt_providers\\._cryptography",
)

import pdfplumber
try:
    import pytesseract
except ImportError:
    pytesseract = None
from reportlab.pdfgen import canvas
from reportlab.lib.colors import black
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase._fontdata import standardFonts
from reportlab.lib.utils import ImageReader

# 使用绝对导入，避免在脚本直接运行时出现"attempted relative import"错误
from anonymizer_core import anonymize_text


def _guess_font_name(fontname: str) -> str:
    """将 PDF 中的字体映射到 reportlab 内置字体，以最大程度保留样式。"""

    font_upper = (fontname or "").upper()
    if "BOLD" in font_upper:
        return "Helvetica-Bold"
    if "ITALIC" in font_upper or "OBLIQUE" in font_upper:
        return "Helvetica-Oblique"
    return "Helvetica"


def _normalize_font(fontname: str) -> str:
    """返回一个 reportlab 已注册的字体名称，避免绘制时报错。"""

    guessed = _guess_font_name(fontname)
    if guessed in pdfmetrics.getRegisteredFontNames():
        return guessed
    return "Helvetica"


def _get_text_width(text: str, font_name: str, font_size: float) -> float:
    """计算文本在给定字体和字号下的宽度（单位：点）。"""

    if not text:
        return 0.0

    try:
        font = pdfmetrics.getFont(font_name)
        width = 0.0
        for char in text:
            width += font.face.charWidths.get(ord(char), 500)
        return width * font_size / 1000.0
    except:
        # 如果无法获取字体信息，使用估算值（平均字符宽度）
        return len(text) * font_size * 0.6


def _fit_text_to_width(text: str, font_name: str, font_size: float, max_width: float, min_font_size: float = 4.0) -> tuple:
    """
    调整文本以适应给定宽度。

    返回: (adjusted_text, adjusted_font_size)
    - 如果文本适合，返回原始文本和字号
    - 如果不适合，尝试缩小字号
    - 如果缩小到最小字号仍不适合，截断文本
    """

    if not text or max_width <= 0:
        return text, font_size

    # 计算当前宽度
    current_width = _get_text_width(text, font_name, font_size)

    # 如果适合，直接返回
    if current_width <= max_width:
        return text, font_size

    # 尝试缩小字号
    adjusted_size = font_size * (max_width / current_width)
    if adjusted_size >= min_font_size:
        return text, adjusted_size

    # 如果缩小到最小字号仍不适合，使用最小字号并截断文本
    adjusted_size = min_font_size
    chars_that_fit = int(len(text) * max_width / current_width)
    if chars_that_fit < len(text):
        return text[:max(1, chars_that_fit - 3)] + "...", adjusted_size

    return text, adjusted_size


def _is_tesseract_available() -> bool:
    return shutil.which("tesseract") is not None


def _ocr_page_to_words(image, dpi: int):
    """使用 Tesseract OCR 将单页图像解析为带坐标的词列表。"""

    data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
    scale = 72.0 / dpi  # 将像素转换为 PDF 点

    words = []
    for i, text in enumerate(data.get("text", [])):
        if not text or text.isspace():
            continue

        conf = int(data.get("conf", ["0"])[i])
        if conf < 0:  # 丢弃无效识别
            continue

        left = data["left"][i] * scale
        top = data["top"][i] * scale
        width = data["width"][i] * scale
        height = data["height"][i] * scale

        words.append(
            {
                "text": text,
                "x0": left,
                "x1": left + width,
                "top": top,
                "bottom": top + height,
                "width": width,
                "height": height,
                "size": height,  # 近似字号
                "fontname": "Helvetica",
            }
        )

    return words


def pdf_to_xml_with_ocr(input_path: str, dpi: int = 200):
    """
    将 PDF 转换为 XML，同时保留页面背景，便于在匿名化后还原排版。

    - 使用 pdfplumber + pypdfium2 将页面渲染为高分辨率图像。
    - 使用 Tesseract OCR 提取带位置信息的文字。
    - 返回 XML 根节点和每页的背景图像列表。
    """

    if not _is_tesseract_available():
        raise RuntimeError("Tesseract OCR 未安装，无法执行 OCR 流程")

    root = ET.Element("document")
    page_images = []

    with pdfplumber.open(input_path) as pdf:
        for page_index, page in enumerate(pdf.pages):
            pil_image = page.to_image(resolution=dpi).original
            page_images.append(pil_image)

            width_px, height_px = pil_image.size
            width_pt = width_px * 72.0 / dpi
            height_pt = height_px * 72.0 / dpi

            page_el = ET.SubElement(
                root,
                "page",
                index=str(page_index),
                width=str(width_pt),
                height=str(height_pt),
            )

            ocr_words = _ocr_page_to_words(pil_image, dpi=dpi)
            ocr_words.sort(key=lambda w: (round(w.get("top", 0), 1), w.get("x0", 0)))

            current_top = None
            line_el = None
            for word in ocr_words:
                rounded_top = round(word.get("top", 0), 1)
                tolerance = 2.0

                if current_top is None or abs(rounded_top - current_top) > tolerance:
                    line_el = ET.SubElement(page_el, "line", top=str(rounded_top))
                    current_top = rounded_top

                ET.SubElement(
                    line_el,
                    "word",
                    x0=str(word.get("x0", 0)),
                    x1=str(word.get("x1", 0)),
                    top=str(word.get("top", 0)),
                    bottom=str(word.get("bottom", 0)),
                    width=str(word.get("width", 0)),
                    height=str(word.get("height", 0)),
                    font=_guess_font_name(word.get("fontname", "")),
                    size=str(word.get("size", 10)),
                    upright="True",
                ).text = word.get("text", "")

    return root, page_images


def pdf_to_xml(input_path: str) -> ET.Element:
    """使用 pdfplumber 将 PDF 转为包含位置和字体信息的 XML。"""

    root = ET.Element("document")
    with pdfplumber.open(input_path) as pdf:
        for page_index, page in enumerate(pdf.pages):
            page_el = ET.SubElement(
                root,
                "page",
                index=str(page_index),
                width=str(page.width),
                height=str(page.height),
            )

            words = page.extract_words(
                extra_attrs=["fontname", "size"],
            )

            # 按行聚合，确保样式和位置更贴近原文
            # 使用更精确的排序和分组逻辑
            words.sort(key=lambda w: (round(w.get("top", 0), 1), w.get("x0", 0)))
            current_top = None
            line_el = None
            for word in words:
                word_top = word.get("top", 0)
                word_size = word.get("size", 10) or 10
                rounded_top = round(word_top, 1)

                # 使用固定容差值（2个点），更稳定
                tolerance = 2.0

                if current_top is None or abs(rounded_top - current_top) > tolerance:
                    line_el = ET.SubElement(page_el, "line", top=str(rounded_top))
                    current_top = rounded_top

                # 保存单词的右边界，用于后续的宽度检测
                word_x0 = word.get("x0", 0)
                word_x1 = word.get("x1", 0)
                word_width = word_x1 - word_x0
                word_bottom = word.get("bottom", 0)
                word_height = word_bottom - word_top

                ET.SubElement(
                    line_el,
                    "word",
                    x0=str(word_x0),
                    x1=str(word_x1),
                    top=str(word_top),
                    bottom=str(word_bottom),
                    width=str(word_width),
                    height=str(word_height),
                    font=_guess_font_name(word.get("fontname", "")),
                    size=str(word_size),
                    upright=str(word.get("upright", True)),
                ).text = word.get("text", "")

    return root


def anonymize_xml(xml_root: ET.Element, config_path: str) -> ET.Element:
    """对 XML 中的每一行文本执行匿名化，同时保留样式节点。"""

    for line_el in xml_root.iter("line"):
        words = list(line_el.iter("word"))
        if not words:
            continue

        original_line = " ".join(word.text or "" for word in words)
        new_line, _ = anonymize_text(original_line, config_path)

        new_tokens = [t for t in new_line.split(" ") if t]  # 过滤空token

        # 改进的Token重分配策略
        if len(new_tokens) == len(words):
            # 精确匹配，直接替换
            for idx in range(len(words)):
                words[idx].text = new_tokens[idx]

        elif len(new_tokens) < len(words):
            # token减少：分配已有的token，剩余单词设为空或保留空格
            for idx in range(len(new_tokens)):
                words[idx].text = new_tokens[idx]
            # 剩余的单词清空，以保持原有的间距
            for idx in range(len(new_tokens), len(words)):
                words[idx].text = ""

        else:
            # token增加：尽量均匀分配，避免全部堆积到最后一个词
            # 策略：将多余的token均匀分配到可用空间较大的单词上
            tokens_per_word = len(new_tokens) // len(words)
            remainder = len(new_tokens) % len(words)

            token_idx = 0
            for word_idx in range(len(words)):
                # 基础分配
                tokens_to_assign = tokens_per_word
                # 将余数分配给前面的单词
                if word_idx < remainder:
                    tokens_to_assign += 1

                # 合并多个token到一个单词
                assigned_tokens = new_tokens[token_idx:token_idx + tokens_to_assign]
                words[word_idx].text = " ".join(assigned_tokens)
                token_idx += tokens_to_assign

    return xml_root


def xml_to_pdf(xml_root: ET.Element, output_path: str, backgrounds=None):
    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    buffer = BytesIO()
    c = canvas.Canvas(buffer)
    c.setFillColor(black)

    for page_index, page_el in enumerate(xml_root.findall("page")):
        width = float(page_el.get("width", 595.2))
        height = float(page_el.get("height", 841.8))
        c.setPageSize((width, height))

        if backgrounds and page_index < len(backgrounds):
            bg_image = ImageReader(backgrounds[page_index])
            c.drawImage(bg_image, 0, 0, width=width, height=height)

        for line_el in page_el.findall("line"):
            for word_el in line_el.findall("word"):
                x0 = float(word_el.get("x0", 40))
                x1 = float(word_el.get("x1", x0 + 100))  # 使用x1作为右边界
                top = float(word_el.get("top", 40))
                size = float(word_el.get("size", 10))
                font = _normalize_font(word_el.get("font", "Helvetica"))
                height_word = float(word_el.get("height", size))

                # 获取文本内容并清理
                text = (word_el.text or "").replace("\n", " ").strip()
                if not text:
                    continue

                # 计算可用宽度（考虑右边界和页面宽度）
                available_width = min(x1 - x0, width - x0 - 10)  # 留10点右边距

                # 如果没有明确的x1或宽度信息，使用默认的较大宽度
                if available_width <= 0 or (x1 - x0) < 1:
                    available_width = width - x0 - 10

                # 调整文本以适应可用宽度
                adjusted_text, adjusted_size = _fit_text_to_width(
                    text, font, size, available_width
                )

                # 改进的基线偏移计算
                # 使用字体的上升高度（ascent）来计算更精确的基线位置
                # 一般字体的基线位置约为高度的 70-75%
                baseline_offset = height_word * 0.75 if height_word else adjusted_size * 0.75

                # pdfplumber 的 top 以页面上边为 0；reportlab 原点在左下
                y = height - top - baseline_offset

                # 绘制文本
                c.setFont(font, adjusted_size)
                c.drawString(x0, y, adjusted_text[:1000])

        c.showPage()

    c.save()
    with open(out_path, "wb") as f:
        f.write(buffer.getvalue())


def anonymize_pdf(input_path: str, output_path: str, config_path: str, use_ocr: bool = False, use_word_pipeline: bool = True):
    """
    匿名化 PDF。

    Args:
        input_path: 输入 PDF 文件路径
        output_path: 输出 PDF 文件路径
        config_path: 脱敏配置文件路径
        use_ocr: 是否使用 OCR 流程（仅在 use_word_pipeline=False 时有效）
        use_word_pipeline: 是否使用 PDF->Word->PDF 流程（推荐，格式保留更好）
    """

    # 优先使用 PyMuPDF 流程，格式保留更完整且更可靠
    if use_word_pipeline:
        try:
            anonymize_pdf_via_pymupdf(input_path, output_path, config_path)
            return
        except Exception as exc:
            warnings.warn(
                f"PyMuPDF 流程不可用，回退到常规解析。原因: {exc}", RuntimeWarning
            )

    # 回退到原有的 XML 流程
    if use_ocr:
        try:
            xml_root, backgrounds = pdf_to_xml_with_ocr(input_path)
        except Exception as exc:
            warnings.warn(
                f"OCR 流程不可用，回退到常规解析。原因: {exc}", RuntimeWarning
            )
            xml_root = pdf_to_xml(input_path)
            backgrounds = None
    else:
        xml_root = pdf_to_xml(input_path)
        backgrounds = None

    anonymized_xml = anonymize_xml(xml_root, config_path)
    xml_to_pdf(anonymized_xml, output_path, backgrounds=backgrounds)


# ============================================================================
# 新的 PDF -> Word -> PDF 流程（保留完整格式）
# ============================================================================

def pdf_to_word(pdf_path: str, docx_path: str):
    """
    使用 pdf2docx 将 PDF 转换为 Word，保留完整的格式和排版。

    Args:
        pdf_path: 输入 PDF 文件路径
        docx_path: 输出 Word 文件路径
    """
    from pdf2docx import Converter

    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
    except Exception as e:
        raise RuntimeError(f"PDF 转 Word 失败: {e}")


def anonymize_word_document(docx_path: str, output_path: str, config_path: str):
    """
    对 Word 文档进行脱敏处理，保留完整的格式和样式。

    Args:
        docx_path: 输入 Word 文件路径
        output_path: 输出 Word 文件路径
        config_path: 脱敏配置文件路径
    """
    from docx import Document

    try:
        doc = Document(docx_path)

        # 处理段落中的文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # 对段落文本进行脱敏
                anonymized_text, _ = anonymize_text(paragraph.text, config_path)

                # 保留格式的文本替换
                _replace_paragraph_text(paragraph, anonymized_text)

        # 处理表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            anonymized_text, _ = anonymize_text(paragraph.text, config_path)
                            _replace_paragraph_text(paragraph, anonymized_text)

        # 处理文本框和形状中的文本（如果有）
        for shape in doc.inline_shapes:
            if hasattr(shape, 'text_frame'):
                for paragraph in shape.text_frame.paragraphs:
                    if paragraph.text.strip():
                        anonymized_text, _ = anonymize_text(paragraph.text, config_path)
                        _replace_paragraph_text(paragraph, anonymized_text)

        # 保存修改后的文档
        doc.save(output_path)

    except Exception as e:
        raise RuntimeError(f"Word 文档脱敏失败: {e}")


def _replace_paragraph_text(paragraph, new_text):
    """
    替换段落文本，同时保留格式（字体、大小、颜色等）。

    Args:
        paragraph: python-docx 段落对象
        new_text: 新的文本内容
    """
    # 清空段落内容但保留格式
    if len(paragraph.runs) > 0:
        # 保留第一个 run 的格式
        first_run = paragraph.runs[0]
        # 清空所有 runs
        for _ in range(len(paragraph.runs)):
            paragraph._element.remove(paragraph.runs[0]._element)
        # 添加新文本，使用第一个 run 的格式
        new_run = paragraph.add_run(new_text)
        # 复制格式
        if first_run.font:
            new_run.font.name = first_run.font.name
            new_run.font.size = first_run.font.size
            new_run.font.bold = first_run.font.bold
            new_run.font.italic = first_run.font.italic
            new_run.font.underline = first_run.font.underline
            new_run.font.color.rgb = first_run.font.color.rgb if first_run.font.color.rgb else None
    else:
        # 如果没有 runs，直接添加文本
        paragraph.add_run(new_text)


def word_to_pdf(docx_path: str, pdf_path: str):
    """
    使用 LibreOffice 将 Word 文档转换为 PDF，保留完整格式。

    Args:
        docx_path: 输入 Word 文件路径
        pdf_path: 输出 PDF 文件路径
    """
    import subprocess
    from pathlib import Path
    import os

    try:
        docx_path_obj = Path(docx_path).resolve()
        pdf_path_obj = Path(pdf_path).resolve()

        # 确保输出目录存在
        pdf_path_obj.parent.mkdir(parents=True, exist_ok=True)

        # LibreOffice 会在与输入文件相同的目录或指定的输出目录中生成 PDF
        # 我们使用临时目录作为输出目录
        temp_output_dir = docx_path_obj.parent

        # 使用 LibreOffice 命令行转换
        # --headless: 无界面模式
        # --convert-to pdf: 转换为 PDF
        # --outdir: 输出目录
        cmd = [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(temp_output_dir),
            str(docx_path_obj),
        ]

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60,
        )

        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice 转换失败:\nstdout: {result.stdout}\nstderr: {result.stderr}")

        # LibreOffice 会生成与输入文件同名的 PDF（但扩展名为 .pdf）
        expected_output = temp_output_dir / (docx_path_obj.stem + ".pdf")

        if not expected_output.exists():
            # 列出目录中的所有文件以便调试
            import os
            files_in_dir = list(os.listdir(temp_output_dir))
            raise RuntimeError(
                f"转换后的 PDF 文件未生成: {expected_output}\n"
                f"目录中的文件: {files_in_dir}\n"
                f"LibreOffice stdout: {result.stdout}\n"
                f"LibreOffice stderr: {result.stderr}"
            )

        # 移动到目标位置
        if expected_output != pdf_path_obj:
            import shutil
            shutil.move(str(expected_output), str(pdf_path_obj))

    except subprocess.TimeoutExpired:
        raise RuntimeError("Word 转 PDF 超时")
    except Exception as e:
        raise RuntimeError(f"Word 转 PDF 失败: {e}")


def anonymize_pdf_via_pymupdf(input_path: str, output_path: str, config_path: str):
    """
    使用 PyMuPDF 直接在 PDF 上进行文本替换，完整保留原始格式。

    这种方法比 PDF->Word->PDF 更可靠，且完全保留原始 PDF 的所有格式、字体和布局。

    Args:
        input_path: 输入 PDF 文件路径
        output_path: 输出 PDF 文件路径
        config_path: 脱敏配置文件路径
    """
    import fitz  # PyMuPDF
    from pathlib import Path

    try:
        # 打开 PDF 文档
        doc = fitz.open(input_path)

        # 遍历每一页
        for page_num in range(len(doc)):
            page = doc[page_num]

            # 获取页面上的所有文本块
            text_dict = page.get_text("dict")
            blocks = text_dict.get("blocks", [])

            for block in blocks:
                if block.get("type") == 0:  # 文本块
                    for line in block.get("lines", []):
                        # 构建行文本
                        line_text = "".join([span.get("text", "") for span in line.get("spans", [])])

                        if not line_text.strip():
                            continue

                        # 对文本进行脱敏
                        anonymized_text, replacements = anonymize_text(line_text, config_path)

                        # 如果文本发生了变化，进行替换
                        if anonymized_text != line_text and replacements:
                            for original, replacement in replacements:
                                if original and original in line_text:
                                    # 使用 PyMuPDF 的 redact 功能进行替换
                                    # 首先查找要替换的文本
                                    areas = page.search_for(original)

                                    for rect in areas:
                                        # 获取该区域的文本属性
                                        # 我们将用白色矩形覆盖原文本，然后写入新文本
                                        page.add_redact_annot(rect, fill=(1, 1, 1))  # 白色填充

                                    # 应用编辑
                                    page.apply_redactions()

                                    # 在相同位置写入新文本
                                    for rect in areas:
                                        # 获取原始文本的字体和大小
                                        # 使用第一个span的属性作为参考
                                        span = line.get("spans", [{}])[0]
                                        font_size = span.get("size", 12)
                                        font_color = span.get("color", 0)  # 黑色

                                        # 写入新文本
                                        page.insert_text(
                                            (rect.x0, rect.y1 - 2),  # 位置稍微调整
                                            replacement,
                                            fontsize=font_size,
                                            color=_int_to_rgb(font_color) if font_color else (0, 0, 0),
                                        )

        # 保存修改后的 PDF
        doc.save(output_path)
        doc.close()

    except Exception as e:
        raise RuntimeError(f"PDF 脱敏失败: {e}")


def _int_to_rgb(color_int):
    """将整数颜色转换为 RGB 元组"""
    if color_int == 0:
        return (0, 0, 0)  # 黑色
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    return (r / 255.0, g / 255.0, b / 255.0)


def anonymize_pdf_via_word(input_path: str, output_path: str, config_path: str):
    """
    通过 PDF -> Word -> 脱敏 -> PDF 的流程处理 PDF，完整保留格式。

    注意：在某些环境中，LibreOffice 可能无法工作，建议使用 anonymize_pdf_via_pymupdf 替代。

    Args:
        input_path: 输入 PDF 文件路径
        output_path: 输出 PDF 文件路径
        config_path: 脱敏配置文件路径
    """
    import tempfile
    from pathlib import Path

    try:
        # 创建临时目录
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_docx = Path(temp_dir) / "temp.docx"
            temp_anonymized_docx = Path(temp_dir) / "temp_anonymized.docx"

            # 步骤 1: PDF -> Word
            pdf_to_word(input_path, str(temp_docx))

            # 步骤 2: Word 脱敏
            anonymize_word_document(str(temp_docx), str(temp_anonymized_docx), config_path)

            # 步骤 3: Word -> PDF
            word_to_pdf(str(temp_anonymized_docx), output_path)

    except Exception as e:
        raise RuntimeError(f"PDF 脱敏流程失败: {e}")
